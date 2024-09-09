# test.py
import os
import re
import json
import jieba
import pandas as pd
from datetime import datetime
from neo4j import GraphDatabase
from docx import Document
from zhipuai import ZhipuAI
import xml.etree.ElementTree as ET
from flask import current_app
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import sessionmaker
from celery_config import create_app
from langchain.text_splitter import RecursiveCharacterTextSplitter
from models import db, Entity, Relation, EntityReview, RelationshipModel
from neo4j_connector import Neo4jConnector

# 配置信息
API_KEY = os.getenv('API_KEY', 'c0828a92a058f312295e323ba41bccc5.dGOtblQOUy8UeSr8')
# DOCX_PATH = 'process_data/水利大辞典/输出/0 水利 水利史/水利 水利史-水利 - 测试.docx'
DOCX_PATH = 'upload_files/大渡河多业主梯级水电站联合调度机制初探_陈在妮.docx'
# path = 'upload_files/大渡河多业主梯级水电站联合调度机制初探_陈在妮.docx'
OUTPUT_JSON_PATH = 'filtered_prompts.json'
OUTPUT_CSV_PATH = 'filtered_triples.csv'

def initialize_db_session():
    """创建数据库会话"""
    with app.app_context():
        db_session = db.session
    return db_session

# 假设 create_app 是用于创建 Flask 应用的函数
app = create_app()  # 创建 Flask 应用实例
db.init_app(app)  # 初始化数据库

class DocxProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.text = None
        self.paragraphs = None
        self.sentences = None
        self.entities = None

    def read_docx(self):
        doc = Document(self.file_path)
        self.text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
        return self.text

    def split_paragraphs(self):
        if not self.text:
            self.read_docx()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0, separators=["\n\n"])
        self.paragraphs = text_splitter.split_text(self.text)
        return self.paragraphs

    def split_sentences(self, paragraph):
        sentences = re.split(r'(?<=[。！？])', paragraph)
        return [sentence.strip() for sentence in sentences if sentence.strip() != ""]

    def process_sentences(self):
        if not self.paragraphs:
            self.split_paragraphs()
        self.sentences = [sentence for para in self.paragraphs for sentence in self.split_sentences(para)]
        return self.sentences

    def format_for_chatglm(self):
        if not self.sentences:
            self.process_sentences()
        return {f"Prompt_{i+1}": sentence for i, sentence in enumerate(self.sentences)}

    def load_entities_from_db(self):
        try:
            with app.app_context():
                self.entities = db.session.query(Entity).all()
                print("Loaded entities:", [entity.entity_name for entity in self.entities])
        except Exception as e:
            print(f"Error loading entities from DB: {e}")
            self.entities = []
        return self.entities

    def filter_prompts_by_entities(self, prompts):
        if self.entities is None:
            self.load_entities_from_db()

        filtered_prompts = {}
        total_filtered = 0  # 初始化计数器
        patterns = [(re.compile(re.escape(entity.entity_name.lower())), entity) for entity in self.entities]

        for key, content in prompts.items():
            masked_content = content.lower().replace('\n', '').replace(' ', '')

            matched_entities = [
                {"entity_name": entity.entity_name}
                for pattern, entity in patterns for m in pattern.finditer(masked_content)
            ]

            unique_entities = {entity['entity_name']: entity for entity in matched_entities}.values()

            if unique_entities:
                filtered_prompts[key] = {"entity_info": list(unique_entities), "content": content}
                total_filtered += 1  # 计数加一

        return filtered_prompts, total_filtered


def save_json(data, filepath):
    """保存 JSON 数据到文件"""
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


class GLMProcessor:
    def __init__(self, api_key, db_session, model="glm-4-air"):
        self.zhipu = ZhipuAI(api_key=api_key)
        self.model = model
        self.db_session = db_session
        self.all_responses = []
        self.relations = self.load_relations()
        print("Loaded Relations:", self.relations)

    def load_relations(self):
        # 从 relationship_models 表中加载所有关系及其开始和结束节点类型
        try:
            with app.app_context():
                relations = db.session.query(RelationshipModel).all()

                # 打印每个关系的详细信息
                # for r in relations:
                #     print(
                #         f"Relation Name: {r.relation_name}, Start Node Type: {r.start_node_type}, End Node Type: {r.end_node_type}")

                # 将关系信息格式化为字典列表返回
                return [
                    {
                        "relation_name": r.relation_name,
                        "start_node_type": r.start_node_type,
                        "end_node_type": r.end_node_type
                    }
                    for r in relations
                ]
        except Exception as e:
            print(f"Error loading relations from the database: {str(e)}")
            return []


    def extract_triples(self, filtered_prompts, batch_size=3):
        # 构建提示词
        prompt_1 = "你是一名资深的三元组抽取专家，接下来我会给你一些包含实体的文本，你需要从中提取出三元组。"

        # 构建关系的提示词，明确关系名称和限制条件
        relation_prompts = "\n".join(
            [
                f"- 关系: {r['relation_name']}, 起始节点类型: {r['start_node_type']}, 结束节点类型: {r['end_node_type']}"
                for r in self.relations
            ]
        )
        print(relation_prompts)

        prompt_2 = (
            "给你输入的数据格式如下：\n"
            "{\n"
            "    \"Prompt_1\": {\n"
            "        \"entity_info\": [\n"
            "            {\"entity_name\": \"水利\"},\n"
            "            {\"entity_name\": \"水资源\"}\n"
            "        ],\n"
            "        \"content\": \"水利( water conservancy) 人类采取各种措施对 自然界的水进行控制、调节、治导、利用、管理和保 护，以减轻或免除水旱灾害，并开发利用水资源，适 应生产生活和改善生态环境需要的活动。\"\n"
            "    }\n"
            "}\n"
            "1. 你的任务是从这些文本中提取三元组。注意：三元组的关系、起始节点类型和结束节点类型必须从给定的列表中选择，不能自定义。\n"
            "每个三元组的格式应如下：\n"
            "<Triples>\n"
            "    <Triple>\n"
            "        <Head>实体名</Head>\n"
            "        <Relation>关系</Relation>\n"
            "        <Tail>实体名</Tail>\n"
            "        <start_node_Type>开始节点类型</start_node_Type>\n"
            "        <end_node_Type>结束节点类型</end_node_Type>\n"
            "    </Triple>\n"
            "</Triples>\n"
            "2. 你只能从以下关系列表中选择关系、起始节点类型和结束节点类型：\n"
            f"{relation_prompts}\n"
            "3. 在抽取三元组时，必须严格按照起始和结束节点类型的要求。'start_node_Type' 字段用于记录开始节点类型类型，'end_node_Type' 字段用于记录结束节点类型类型，\n"
            "4. 如果无法找到逻辑通顺的三元组，请输出：\n"
            "<Triples>\n"
            "    <Triple>\n"
            "        <Head>无</Head>\n"
            "        <Relation>无</Relation>\n"
            "        <Tail>无</Tail>\n"
            "        <start_node_Type>无</start_node_Type>\n"
            "        <end_node_Type>无</end_node_Type>\n"
            "    </Triple>\n"
            "</Triples>\n"
            "5. 实体说明：其中\"头实体\"只能在\"entity_info\"中选择，不能自定义头实体，但是尾实体则可以从文本中抽取。\n"
            "6. 关系说明：关系必须从给定列表中选择，不得自定义。头实体和尾实体的类型必须与规定的起始和结束节点类型一致。\n"
            "7. 总结：你的回复永远只有XML格式的语句。\n"
            "8. 下面是一个输出的例子：\n"
            "当我输入：\n"
            "\"Prompt_4\": {\n"
            "\"entity_info\": [\n"
            "{\n"
            "\"entity_name\": \"大渡河\"\n"
            "},\n"
            "{\n"
            "\"entity_name\": \"典型河流\"\n"
            "}\n"
            "],\n"
            "\"content\": \"大渡河属于多开发主体的典型河流”。\n"
            "},\n"
            "你的回答应该是类似于：\n"
            "<Triples>\n"
            "    <Triple>\n"
            "        <Head>大渡河</Head>\n"
            "        <Relation>位于</Relation>\n"
            "        <Tail>四川</Tail>\n"
            "        <start_node_Type>河流</start_node_Type>\n"
            "        <end_node_Type>位置</end_node_Type>\n"
            "    </Triple>\n"
            "</Triples>\n"
        )

        prompt_keys = list(filtered_prompts.keys())
        total_prompts = len(prompt_keys)

        for i in range(0, total_prompts, batch_size):
            batch_prompts = {key: filtered_prompts[key] for key in prompt_keys[i:i + batch_size]}
            messages = [
                {"role": "system", "content": prompt_1},
                {"role": "user", "content": prompt_2},
                {"role": "user", "content": json.dumps(batch_prompts, ensure_ascii=False)}
            ]

            response = self.zhipu.chat.completions.create(
                model=self.model,
                messages=messages
            )

            if response:
                response_content = response.choices[0].message.content
                print(response_content)
                self.all_responses.append(response_content)

        return self.all_responses

    # def extract_triples(self, filtered_prompts, batch_size=3):
    #     # 构建提示词
    #     prompt_1 = "你是一名资深的三元组抽取专家，接下来我会给你一些包含实体的文本，你需要从中提取出三元组。"
    #     prompt_2 = (
    #         "给你输入的数据格式如下：\n"
    #         "{\n"
    #         "    \"Prompt_1\": {\n"
    #         "        \"entity_info\": [\n"
    #         "            {\"entity_name\": \"水利\"},\n"
    #         "            {\"entity_name\": \"水资源\"}\n"
    #         "        ],\n"
    #         "        \"content\": \"水利( water conservancy) 人类采取各种措施对 自然界的水进行控制、调节、治导、利用、管理和保 护，以减轻或免除水旱灾害，并开发利用水资源，适 应生产生活和改善生态环境需要的活动。\"\n"
    #         "    }\n"
    #         "}\n"
    #         "1. 你的任务是从这些文本中提取三元组。注意：三元组的关系只能从给定的关系列表中选择"
    #         "不要自行创造关系名称。\n"
    #         "每个三元组的格式应如下：\n"
    #         "<Triples>\n"
    #         "    <Triple>\n"
    #         "        <Head>实体名</Head>\n"
    #         "        <Relation>关系</Relation>\n"
    #         "        <Tail>实体名或属性值</Tail>\n"
    #         "        <Type>关系类型</Type>\n"
    #         "        <Content>原文内容</Content>\n"
    #         "    </Triple>\n"
    #         "</Triples>\n"
    #         "2. 关系列表如下，请根据关系的类型进行选择：\n"
    #         " - 实体关系（用于描述实体与实体之间的关系）：\n"
    #         f"{', '.join([r['relation'] for r in self.relations if r['type'] == '实体关系'])}\n"
    #         " - 属性关系（用于描述实体的属性）：\n"
    #         f"{', '.join([r['relation'] for r in self.relations if r['type'] == '属性关系'])}\n"
    #         "3. 在三元组中，'Type' 字段用于指定关系类型，'实体关系' 或 '属性关系'。\n"
    #         "4. 如果无法找到逻辑通顺的三元组，请输出：\n"
    #         "<Triples>\n"
    #         "    <Triple>\n"
    #         "        <Head>无</Head>\n"
    #         "        <Relation>无</Relation>\n"
    #         "        <Tail>无</Tail>\n"
    #         "        <Type>无</Type>\n"
    #         "        <Content>原文内容</Content>\n"
    #         "    </Triple>\n"
    #         "</Triples>\n"
    #         "5. 实体说明：其中\"头实体\"只能在\"entity_info\"中选择，千万千万不能自定义头实体,这一点请你一定注意，但是尾实体则可以从文本中抽取。"
    #         "例如上述例子中，头实体只能在“水利”，“水资源”当中选择，不能自己定义实体！且头实体的字数没有超过12的。另外，在抽取实体的过程中，如果遇到书名，千万不要加上书名号“《》”！\n"
    #         "6. 关系说明：不得自定义关系，只能从给定的关系列表中选择，一般情况下，一个实体的一个属性关系只存在一个，例如：“水利”的属性关系“英文名”只存在一个“water conservancy”，不存在多个的情况，请一定注意！！！\n"
    #         "7. 总结：你的回复永远只有XML格式的语句。\n"
    #         "8. 下面是一个输出的例子：\n"
    #         "当我输入：\n"
    #         "\"Prompt_4\": {\n"
    #         "\"entity_info\": [\n"
    #         "{\n"
    #         "\"entity_name\": \"土质学\"\n"
    #         "},\n"
    #         "{\n"
    #         "\"entity_name\": \"工程岩土学\"\n"
    #         "}\n"
    #         "],\n"
    #         "\"content\": \"土质学(soilscience)亦称“工程岩土学”。\n"
    #         "},\n"
    #         "你的回答应该是类似于：\n"
    #         "<Triples>\n"
    #         "    <Triple>\n"
    #         "        <Head>土质学</Head>\n"
    #         "        <Relation>亦称</Relation>\n"
    #         "        <Tail>工程岩土学</Tail>\n"
    #         "        <Type>实体关系</Type>\n"
    #         "        <Content>土质学(soilscience)亦称“工程岩土学”。</Content>\n"
    #         "    </Triple>\n"
    #         "</Triples>"
    #     )
    #
    #     prompt_keys = list(filtered_prompts.keys())
    #     total_prompts = len(prompt_keys)
    #
    #     for i in range(0, total_prompts, batch_size):
    #         batch_prompts = {key: filtered_prompts[key] for key in prompt_keys[i:i + batch_size]}
    #         messages = [
    #             {"role": "system", "content": prompt_1},
    #             {"role": "user", "content": prompt_2},
    #             {"role": "user", "content": json.dumps(batch_prompts, ensure_ascii=False)}
    #         ]
    #
    #         response = self.zhipu.chat.completions.create(
    #             model=self.model,
    #             messages=messages
    #         )
    #
    #         if response:
    #             response_content = response.choices[0].message.content
    #             print(response_content)
    #             self.all_responses.append(response_content)
    #
    #     return self.all_responses

    def save_responses_to_file(self, filename):
        # 将所有响应保存到文件
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(self.all_responses, f, ensure_ascii=False, indent=4)


# def extract_triples(data):
#     # 改进后的正则表达式以匹配所有可能的 <Triple> 结构
#     pattern = re.compile(
#         r'<Triple>\s*<Head>(.*?)</Head>\s*<Relation>(.*?)</Relation>\s*<Tail>(.*?)</Tail>\s*<start_node_Type>(.*?)</start_node_Type>\s*<end_node_Type>(.*?)</end_node_Type>\s*</Triple>',
#         re.DOTALL
#     )
#
#     # 用于存储所有匹配的三元组
#     triples = []
#
#     # 遍历每个响应数据
#     for entry in data:
#         # 清理字符串中的多余换行和空格
#         cleaned_entry = entry.replace('\n', '').replace('\r', '').strip()
#
#         # 匹配所有的 <Triple> 结构
#         matches = pattern.findall(cleaned_entry)
#
#         # 如果有匹配的结果，将其添加到 triples 列表中
#         if matches:
#             for match in matches:
#                 head, relation, tail, head_type, tail_type = match
#                 triples.append({
#                     "Head": head.strip(),
#                     "Relation": relation.strip(),
#                     "Tail": tail.strip(),
#                     "Head Type": head_type.strip(),
#                     "Tail Type": tail_type.strip(),
#                     "Tail Major Classification": "",  # 设为空值
#                     "Tail Minor Classification": "",  # 设为空值
#                     "Head Major Classification": "",  # 设为空值
#                     "Head Minor Classification": "",  # 设为空值
#                     "Relation Type": "",  # 设为空值
#                 })
#         else:
#             print(f"No matches found in entry: {cleaned_entry[:100]}...")  # 打印未匹配的部分进行调试
#
#     return triples
#
# def save_triples_to_csv(triples, output_csv_path):
#     # 检查 triples 是否是有效的列表
#     if isinstance(triples, list) and all(isinstance(triple, dict) for triple in triples):
#         try:
#             df = pd.DataFrame(triples)
#             print("DataFrame created successfully, preview:\n", df.head())  # 打印 DataFrame 的头部进行调试
#             df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')
#             print(f"Saved {len(triples)} triples to {output_csv_path}")
#         except Exception as e:
#             print(f"Error saving to CSV: {e}")  # 打印错误信息
#     else:
#         print(
#             f"Error: triples is not a list of dictionaries or is empty. Type of triples: {type(triples)}, value: {triples}")


class TripleProcessor:
    def __init__(self):
        pass

    def extract_triples(self, data):
        """
        从响应数据中提取三元组，并进行过滤。
        :param data: 包含三元组的响应数据列表
        :return: 提取的三元组列表
        """
        # 正则表达式用于匹配 <Triple> 结构
        pattern = re.compile(
            r'<Triple>\s*<Head>(.*?)</Head>\s*<Relation>(.*?)</Relation>\s*<Tail>(.*?)</Tail>\s*<start_node_Type>(.*?)</start_node_Type>\s*<end_node_Type>(.*?)</end_node_Type>\s*</Triple>',
            re.DOTALL
        )

        triples = []

        # 遍历每个响应数据
        for entry in data:
            # 清理字符串中的多余换行和空格
            cleaned_entry = entry.replace('\n', '').replace('\r', '').strip()

            # 匹配所有的 <Triple> 结构
            matches = pattern.findall(cleaned_entry)

            # 如果有匹配的结果，将其添加到 triples 列表中
            if matches:
                for match in matches:
                    head, relation, tail, head_type, tail_type = match
                    triples.append({
                        "Head": head.strip(),
                        "Relation": relation.strip(),
                        "Tail": tail.strip(),
                        "Head Type": head_type.strip(),
                        "Tail Type": tail_type.strip(),
                        "Head Major Classification": "",  # 设为空值
                        "Head Minor Classification": "",  # 设为空值
                        "Tail Major Classification": "",  # 设为空值
                        "Tail Minor Classification": "",  # 设为空值
                        "Relation Type": ""  # 设为空值
                    })
            else:
                print(f"No matches found in entry: {cleaned_entry[:100]}...")  # 打印未匹配的部分进行调试

        return triples

    def save_triples_to_csv(self, triples, output_csv_path):
        """
        将提取的三元组保存为CSV文件。
        :param triples: 提取的三元组列表
        :param output_csv_path: CSV文件保存路径
        """
        # 检查 triples 是否是有效的列表
        if isinstance(triples, list) and all(isinstance(triple, dict) for triple in triples):
            try:
                df = pd.DataFrame(triples)
                print("DataFrame created successfully, preview:\n", df.head())  # 打印 DataFrame 的头部进行调试
                df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')
                print(f"Saved {len(triples)} triples to {output_csv_path}")
            except Exception as e:
                print(f"Error saving to CSV: {e}")  # 打印错误信息
        else:
            print(
                f"Error: triples is not a list of dictionaries or is empty. Type of triples: {type(triples)}, value: {triples}"
            )


class Neo4jSaver:
    def __init__(self, uri, user, password):
        self.driver = GraphDatabase.driver(uri, auth=(user, password))

    def close(self):
        self.driver.close()

    def create_relationship(self, head, relation_name, tail, head_type, tail_type, head_major, head_minor, tail_major,
                            tail_minor, relation_type, docx_id, csv_id):
        head = "" if pd.isna(head) else head
        tail = "" if pd.isna(tail) else tail
        head_type = "" if pd.isna(head_type) else head_type
        tail_type = "" if pd.isna(tail_type) else tail_type
        head_major = "" if pd.isna(head_major) else head_major
        head_minor = "" if pd.isna(head_minor) else head_minor
        tail_major = "" if pd.isna(tail_major) else tail_major
        tail_minor = "" if pd.isna(tail_minor) else tail_minor
        relation_name = "" if pd.isna(relation_name) else relation_name
        relation_type = "" if pd.isna(relation_type) else relation_type
        docx_id = "" if pd.isna(docx_id) else docx_id
        csv_id = "" if pd.isna(csv_id) else csv_id

        with self.driver.session() as session:
            session.execute_write(
                self._create_and_return_relationship, head, relation_name, tail, head_type, tail_type, head_major,
                head_minor, tail_major, tail_minor, relation_type, docx_id, csv_id
            )

    def _create_and_return_relationship(self, tx, head, relation_name, tail, head_type, tail_type, head_major,
                                        head_minor, tail_major, tail_minor, relation_type, docx_id, csv_id):
        created_at = datetime.now().isoformat()
        query = (
            "MERGE (a:Entity {name: $head, type: $head_type, major_classification: $head_major, minor_classification: $head_minor, docx_id: $docx_id, csv_id: $csv_id}) "
            "MERGE (b:Entity {name: $tail, type: $tail_type, major_classification: $tail_major, minor_classification: $tail_minor, docx_id: $docx_id, csv_id: $csv_id}) "
            "MERGE (a)-[r:RELATIONSHIP {name: $relation_name, type: $relation_type, docx_id: $docx_id, csv_id: $csv_id, created_at: $created_at}]->(b) "
            "RETURN a, r, b"
        )

        tx.run(
            query,
            head=head,
            tail=tail,
            head_type=head_type,
            tail_type=tail_type,
            head_major=head_major,
            head_minor=head_minor,
            tail_major=tail_major,
            tail_minor=tail_minor,
            relation_name=relation_name,
            relation_type=relation_type,
            docx_id=docx_id,
            csv_id=csv_id,
            created_at=created_at
        )

    def save_triples_to_neo4j(self, triples_csv_path, docx_id=None, csv_id=None):
        """
        从CSV文件中读取三元组并存储到Neo4j数据库
        :param triples_csv_path: CSV文件路径
        :param docx_id: 文档ID
        :param csv_id: CSV ID
        """
        try:
            triples_df = pd.read_csv(triples_csv_path)
            for _, row in triples_df.iterrows():
                self.create_relationship(
                    row["Head"],
                    row["Relation"],
                    row["Tail"],
                    row["Head Type"],
                    row["Tail Type"],
                    row["Head Major Classification"],
                    row["Head Minor Classification"],
                    row["Tail Major Classification"],
                    row["Tail Minor Classification"],
                    row["Relation Type"],
                    docx_id,
                    csv_id
                )
        except Exception as e:
            print(f"Error saving triples to Neo4j: {e}")


# class TripleRefinementProcessor:
#     def __init__(self, output_folder, docx_name):
#         self.output_folder = output_folder
#         self.docx_name = docx_name
#         self.entities = self.load_entities()
#         self.relations = self.load_relations()
#         self.ensure_directory()
#
#     def ensure_directory(self):
#         self.docx_dir = os.path.join(self.output_folder, self.docx_name)
#         os.makedirs(self.docx_dir, exist_ok=True)
#
#     def load_entities(self):
#         try:
#             with app.app_context():
#                 entities = db.session.query(Entity).all()
#                 print("Loaded entities:", [entity.entity_name for entity in entities])
#                 return [entity.entity_name for entity in entities]
#         except Exception as e:
#             print(f"Error loading entities from the database: {str(e)}")
#             return []
#
#     def load_relations(self):
#         try:
#             with app.app_context():
#                 relations = db.session.query(Relation).all()
#                 print("Loaded relations:", [r.relation for r in relations])
#                 return [relation.relation for relation in relations]
#         except Exception as e:
#             print(f"Error loading relations from the database: {str(e)}")
#             return []
#
#     def parse_triples(self, xml_response):
#         xml_fragments = xml_response.strip().split('<?xml version="1.0"?>')
#         triples = []
#         for fragment in xml_fragments:
#             fragment = fragment.strip()
#             if not fragment:
#                 continue
#             try:
#                 if not fragment.startswith('<Triples>'):
#                     fragment = '<Triples>' + fragment
#                 if not fragment.endswith('</Triples>'):
#                     fragment += '</Triples>'
#
#                 root = ET.fromstring(fragment)
#                 for triple in root.findall(".//Triple"):
#                     head_elem = triple.find("Head")
#                     relation_elem = triple.find("Relation")
#                     tail_elem = triple.find("Tail")
#                     type_elem = triple.find("Type")
#                     content_elem = triple.find("Content")
#
#                     head = head_elem.text if head_elem is not None else ""
#                     relation = relation_elem.text if relation_elem is not None else ""
#                     tail = tail_elem.text if tail_elem is not None else ""
#                     type_ = type_elem.text if type_elem is not None else ""
#                     content = content_elem.text if content_elem is not None else ""
#
#                     triples.append({
#                         "Head": head,
#                         "Relation": relation,
#                         "Tail": tail,
#                         "Type": type_,
#                         "Content": content
#                     })
#             except ET.ParseError as e:
#                 print(f"Error parsing XML fragment: {e}")
#                 print(f"Problematic fragment: {fragment}")
#
#         return triples
#
#     def refine_triples(self, all_responses):
#         all_triples = []
#         valid_triples = []
#         invalid_triples = []
#
#         for response in all_responses:
#             triples = self.parse_triples(response)
#             all_triples.extend(triples)
#             for triple in triples:
#                 if triple["Head"] in self.entities and triple["Relation"] in self.relations:
#                     valid_triples.append(triple)
#                 else:
#                     invalid_triples.append(triple)
#
#         self.save_to_csv(all_triples, "all_triples.csv")
#         self.save_to_csv(valid_triples, "valid_triples.csv")
#         self.save_to_csv(invalid_triples, "invalid_triples.csv")
#
#         # 将有效的三元组保存到 MySQL 的 EntityReview 表中
#         self.save_to_entity_review(valid_triples)
#
#     def save_to_csv(self, data, filename):
#         df = pd.DataFrame(data)
#         df.to_csv(os.path.join(self.docx_dir, filename), index=False, encoding='utf-8-sig')
#
#     def save_to_entity_review(self, data):
#         try:
#             with app.app_context():
#                 # 清空表中的旧数据
#                 db.session.query(EntityReview).delete()
#                 db.session.commit()
#
#                 # 插入新数据
#                 for triple in data:
#                     # 保存 Head 实体
#                     head_review_entry = EntityReview(
#                         entity_name=triple["Head"],
#                         source_file=self.docx_name,
#                         source_text=triple["Content"],
#                         review_status="待审核",
#                         review_action="",
#                         merge_to_entity=""
#                     )
#                     db.session.add(head_review_entry)
#
#                     # 如果 Tail 也是一个实体并且需要审核
#                     if triple["Type"] == "实体关系":  # 假设 Type 表示实体关系
#                         tail_review_entry = EntityReview(
#                             entity_name=triple["Tail"],
#                             source_file=self.docx_name,
#                             source_text=triple["Content"],
#                             review_status="待审核",
#                             review_action="",
#                             merge_to_entity=""
#                         )
#                         db.session.add(tail_review_entry)
#
#                 db.session.commit()
#                 print("New data saved to entity_reviews.")
#         except Exception as e:
#             db.session.rollback()  # 出现异常时回滚事务
#             print(f"Error saving to entity_reviews: {str(e)}")
#
#
#     def update_neo4j_from_review(self, neo4j_uri, neo4j_user, neo4j_password):
#         """将审核通过的三元组从 EntityReview 表同步到 Neo4j"""
#         try:
#             with app.app_context():
#                 # 查询审核通过的条目
#                 approved_entries = EntityReview.query.filter_by(review_status="已审核").all()
#
#                 neo4j_connector = Neo4jConnector(neo4j_uri, neo4j_user, neo4j_password)
#
#                 for entry in approved_entries:
#                     # 属性关系处理
#                     if entry.type == '属性关系':
#                         neo4j_connector.add_property_to_entity(entry.entity_name, entry.relation, entry.tail)
#                     else:
#                         # 实体关系处理
#                         neo4j_connector.create_relationship2(entry.entity_name, entry.relation, entry.tail)
#
#                 neo4j_connector.close()
#
#                 print(f"Approved entries have been saved to Neo4j.")
#
#         except Exception as e:
#             print(f"Error updating Neo4j from EntityReview: {str(e)}")
#
#
#     # Neo4j保存方法保持不变，确保只有审核通过的数据才会同步到 Neo4j
#     def save_to_neo4j(self, csv_file_path, neo4j_uri, neo4j_user, neo4j_password):
#         neo4j_connector = Neo4jConnector(neo4j_uri, neo4j_user, neo4j_password)
#         neo4j_connector.save_to_neo4j(csv_file_path)  # 使用 Neo4jConnector 类的方法进行保存
#         neo4j_connector.close()


def main():
    # 初始化 Flask 应用
    app = create_app()

    # 进入应用上下文
    with app.app_context():
        # 初始化数据库会话
        db_session = initialize_db_session()

        # 文档处理
        processor = DocxProcessor(DOCX_PATH)
        formatted_prompts = processor.format_for_chatglm()

        # 过滤并处理提示
        filtered_prompts, total_filtered = processor.filter_prompts_by_entities(formatted_prompts)

        # 保存过滤后的 JSON 数据
        save_json(filtered_prompts, OUTPUT_JSON_PATH)

        # GLM 三元组提取
        glm_processor = GLMProcessor(API_KEY, db_session)
        glm_responses = glm_processor.extract_triples(filtered_prompts)  # 传入解包后的 filtered_prompts

        # 保存 GLM 的响应到 JSON 文件
        glm_responses_file = 'generate_data/test/glm_responses.json'
        glm_processor.save_responses_to_file(glm_responses_file)

        # 读取之前生成的 GLM 响应 JSON 文件
        # glm_responses_file = 'generate_data/test/glm_responses.json'
        # with open(glm_responses_file, 'r', encoding='utf-8') as f:
        #     data = json.load(f)

        # 初始化 TripleProcessor
        triple_processor = TripleProcessor()

        # 提取三元组
        triples = triple_processor.extract_triples(glm_responses)

        # 打印提取的三元组数量
        print(f"Number of extracted triples: {len(triples)}")

        # CSV 文件保存路径
        output_csv_path = 'extracted_triples.csv'

        # 将三元组保存为 CSV
        triple_processor.save_triples_to_csv(triples, output_csv_path)

        # 连接到 Neo4j 并保存三元组
        neo4j_saver = Neo4jSaver("bolt://localhost:7687", "neo4j", "12345678")
        neo4j_saver.save_triples_to_neo4j(output_csv_path, docx_id="docx123", csv_id="csv123")
        neo4j_saver.close()


if __name__ == "__main__":
    main()

